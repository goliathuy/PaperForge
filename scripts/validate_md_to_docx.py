import sys
import os
import re
from docx import Document
from generate_word_from_md import generate_word_from_md
from test_docx_format import test_docx_format

def validate_md_to_docx(md_file_path, expected_format_rules=None):
    """
    Validates the conversion from Markdown to DOCX by:
    1. Converting the Markdown file to DOCX
    2. Testing the DOCX formatting
    3. Validating against expected formatting rules (if provided)
    
    Args:
        md_file_path (str): Path to the Markdown file
        expected_format_rules (dict, optional): Dictionary containing expected format rules
            Example: {
                'headings': [('Heading 1', 1), ('Heading 2', 2)],
                'lists': ['List Bullet', 'List Number'],
                'styles': ['Normal', 'CodeBlock'],
                'bold_text': ['important phrase']
            }
    
    Returns:
        tuple: (bool, dict) - (validation_passed, validation_results)
    """
    # Create output docx file path
    base_name = os.path.splitext(os.path.basename(md_file_path))[0]
    output_docx_path = os.path.join(os.path.dirname(md_file_path), f"{base_name}_validated.docx")
    
    print(f"Step 1: Converting {md_file_path} to {output_docx_path}")
    
    # Generate the Word document from Markdown
    try:
        generate_word_from_md(md_file_path, output_docx_path)
        print(f"✓ Successfully generated Word document: {output_docx_path}")
    except Exception as e:
        print(f"✗ Failed to generate Word document: {str(e)}")
        return False, {"error": str(e)}
    
    print(f"\nStep 2: Testing DOCX formatting of {output_docx_path}")
    
    # Collect formatting information
    doc = Document(output_docx_path)
    
    # Create a dictionary to store formatting information
    format_info = {
        "headings": [],
        "lists": [],
        "styles": set(),
        "bold_text": [],
        "tables": [],
        "links": [],
        "code_blocks": []
    }
    
    # Analyze document structure
    for i, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name
        text = paragraph.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
            
        # Store style information
        format_info["styles"].add(style)
        
        # Check for headings
        if style.startswith('Heading'):
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 0
            format_info["headings"].append((text, level))
            
        # Check for lists
        if style in ['List Bullet', 'List Number', 'List Bullet 2']:
            format_info["lists"].append(style)
            
        # Check for bold text
        if paragraph.runs:
            for run in paragraph.runs:
                if run.bold and run.text.strip():
                    format_info["bold_text"].append(run.text.strip())
                    
        # Check for code blocks
        if style == 'CodeBlock':
            format_info["code_blocks"].append(text)
    
    # Collect table information
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        format_info["tables"].append(table_data)
    
    # Convert set to list for JSON serialization
    format_info["styles"] = list(format_info["styles"])
    
    # Display formatting information
    print("\nDocument formatting summary:")
    print(f"- Headings: {len(format_info['headings'])}")
    if format_info['headings']:
        for heading, level in format_info['headings'][:3]:
            print(f"  * Level {level}: {heading[:30]}{'...' if len(heading) > 30 else ''}")
        if len(format_info['headings']) > 3:
            print(f"  * ... and {len(format_info['headings']) - 3} more")
            
    print(f"- Lists: {len(format_info['lists'])}")
    list_types = {}
    for list_type in format_info['lists']:
        list_types[list_type] = list_types.get(list_type, 0) + 1
    for list_type, count in list_types.items():
        print(f"  * {list_type}: {count}")
        
    print(f"- Tables: {len(format_info['tables'])}")
    for i, table in enumerate(format_info['tables'][:2]):
        print(f"  * Table {i+1}: {len(table)} rows x {len(table[0]) if table else 0} columns")
    if len(format_info['tables']) > 2:
        print(f"  * ... and {len(format_info['tables']) - 2} more")
        
    print(f"- Code blocks: {len(format_info['code_blocks'])}")
    print(f"- Bold text elements: {len(format_info['bold_text'])}")
    print(f"- Styles used: {', '.join(format_info['styles'])}")
    
    # Validate against expected rules if provided
    validation_results = {"passed": True, "details": {}}
    
    if expected_format_rules:
        print("\nStep 3: Validating against expected format rules")
        
        # Validate headings
        if 'headings' in expected_format_rules:
            expected_headings = expected_format_rules['headings']
            actual_headings = format_info['headings']
            headings_match = True
            
            # Check if we have at least the expected headings
            if len(actual_headings) < len(expected_headings):
                headings_match = False
                validation_results["details"]["headings"] = f"Missing headings: expected {len(expected_headings)}, found {len(actual_headings)}"
            else:
                # Check each expected heading
                for i, (expected_text, expected_level) in enumerate(expected_headings):
                    if i < len(actual_headings):
                        actual_text, actual_level = actual_headings[i]
                        if expected_level != actual_level or not re.search(expected_text, actual_text, re.IGNORECASE):
                            headings_match = False
                            validation_results["details"]["headings"] = f"Heading mismatch at position {i+1}"
                            break
            
            validation_results["passed"] &= headings_match
            print(f"✓ Headings validation: {'Passed' if headings_match else 'Failed'}")
            
        # Validate lists
        if 'lists' in expected_format_rules:
            expected_lists = expected_format_rules['lists']
            actual_lists = format_info['lists']
            
            # Check if all expected list types are present
            lists_match = all(list_type in actual_lists for list_type in expected_lists)
            
            validation_results["passed"] &= lists_match
            validation_results["details"]["lists"] = "All expected list types found" if lists_match else "Missing some expected list types"
            print(f"✓ Lists validation: {'Passed' if lists_match else 'Failed'}")
            
        # Validate styles
        if 'styles' in expected_format_rules:
            expected_styles = expected_format_rules['styles']
            actual_styles = format_info['styles']
            
            # Check if all expected styles are present
            styles_match = all(style in actual_styles for style in expected_styles)
            
            validation_results["passed"] &= styles_match
            validation_results["details"]["styles"] = "All expected styles found" if styles_match else "Missing some expected styles"
            print(f"✓ Styles validation: {'Passed' if styles_match else 'Failed'}")
            
        # Validate bold text
        if 'bold_text' in expected_format_rules:
            expected_bold = expected_format_rules['bold_text']
            actual_bold = format_info['bold_text']
            
            # Check if all expected bold text elements are present
            bold_match = all(any(re.search(expected, actual, re.IGNORECASE) for actual in actual_bold) 
                            for expected in expected_bold)
            
            validation_results["passed"] &= bold_match
            validation_results["details"]["bold_text"] = "All expected bold text found" if bold_match else "Missing some expected bold text"
            print(f"✓ Bold text validation: {'Passed' if bold_match else 'Failed'}")
            
        print(f"\nOverall validation: {'Passed' if validation_results['passed'] else 'Failed'}")
    
    # Run the test_docx_format function to get detailed formatting information
    print("\nDetailed formatting analysis:")
    test_docx_format(output_docx_path)
    
    return validation_results["passed"], format_info

def create_expected_rules_from_md(md_file_path):
    """
    Analyzes a markdown file and creates expected formatting rules
    
    Args:
        md_file_path (str): Path to the Markdown file
        
    Returns:
        dict: Expected format rules
    """
    with open(md_file_path, 'r', encoding='utf-8') as md_file:
        content = md_file.read()
    
    expected_rules = {
        'headings': [],
        'lists': [],
        'styles': ['Normal'],
        'bold_text': []
    }
    
    # Find headings (bold lines that are not within lists)
    heading_pattern = r'^(\*\*.*?\*\*)$'
    headings = re.findall(heading_pattern, content, re.MULTILINE)
    
    # First heading is level 1, rest are level 2
    for i, heading in enumerate(headings):
        text = heading.strip('*')
        level = 1 if i == 0 else 2
        expected_rules['headings'].append((text, level))
    
    # Check for list types
    if re.search(r'^\* ', content, re.MULTILINE):
        expected_rules['lists'].append('List Bullet')
    
    if re.search(r'^\d+\. ', content, re.MULTILINE):
        expected_rules['lists'].append('List Number')
        
    if re.search(r'^    \* ', content, re.MULTILINE):
        expected_rules['lists'].append('List Bullet 2')
    
    # Check for code blocks
    if '```' in content:
        expected_rules['styles'].append('CodeBlock')
    
    # Find bold text
    bold_pattern = r'\*\*(.*?)\*\*'
    bold_matches = re.findall(bold_pattern, content)
    expected_rules['bold_text'] = [match for match in bold_matches if match.strip()]
    
    return expected_rules

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python validate_md_to_docx.py <input_md_file> [--analyze-only]")
        sys.exit(1)
        
    md_file_path = sys.argv[1]
    analyze_only = "--analyze-only" in sys.argv
    
    if not os.path.exists(md_file_path):
        print(f"Error: File {md_file_path} does not exist.")
        sys.exit(1)
    
    # Create expected rules from Markdown
    expected_rules = create_expected_rules_from_md(md_file_path)
    
    if analyze_only:
        print("Markdown analysis results (expected rules):")
        for category, rules in expected_rules.items():
            print(f"{category}: {rules}")
    else:
        # Validate the conversion
        passed, results = validate_md_to_docx(md_file_path, expected_rules)
        
        # Print final result
        print("\n" + "="*50)
        print(f"VALIDATION {'PASSED' if passed else 'FAILED'}")
        print("="*50)