from docx import Document
from docx.enum.style import WD_STYLE_TYPE

def test_docx_format(docx_file_path):
    # Open the Word document
    doc = Document(docx_file_path)

    # Iterate through paragraphs and print their format
    for i, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name
        text = paragraph.text.strip()
        
        # Check for bold, italic, underline
        is_bold = paragraph.runs and any(run.bold for run in paragraph.runs)
        is_italic = paragraph.runs and any(run.italic for run in paragraph.runs)
        is_underlined = paragraph.runs and any(run.underline for run in paragraph.runs)
        
        # Check for color
        color = None
        if paragraph.runs and hasattr(paragraph.runs[0].font.color, 'rgb'):
            color = paragraph.runs[0].font.color.rgb
        
        # Check for list type
        is_list = style in ['List Bullet', 'List Number', 'List Bullet 2']
        list_type = None
        if is_list:
            if style == 'List Bullet':
                list_type = 'Bulleted'
            elif style == 'List Number':
                list_type = 'Numbered'
            elif style == 'List Bullet 2':
                list_type = 'Nested Bulleted'
        
        if text:
            format_info = f"Style: {style}, Bold: {is_bold}, Italic: {is_italic}, Underlined: {is_underlined}, Color: {color}"
            if is_list:
                format_info += f", List Type: {list_type}"
            print(f"Paragraph {i + 1}: '{text}' ({format_info})")
    
    # Iterate through tables and print their contents
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i + 1}:")
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(f"  Row: {row_text}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Usage: python test_docx_format.py <docx_file_path>")
    else:
        docx_file_path = sys.argv[1]
        test_docx_format(docx_file_path)