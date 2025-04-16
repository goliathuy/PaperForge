from docx import Document
import re
import docx.shared
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

def generate_word_from_md(md_file_path, output_file_path):
    # Read the markdown file
    with open(md_file_path, 'r', encoding='utf-8') as md_file:
        lines = md_file.readlines()

    # Create a new Word document
    doc = Document()
    
    # Add a style for code blocks if it doesn't exist
    if 'CodeBlock' not in doc.styles:
        code_block_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_block_style.font.name = 'Courier New'
        code_block_style.font.size = docx.shared.Pt(10)

    # Add a style for List Bullet 2 if it doesn't exist
    if 'List Bullet 2' not in doc.styles:
        list_bullet_2_style = doc.styles.add_style('List Bullet 2', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_2_style.paragraph_format.left_indent = Inches(0.5)
    
    # Add a title style for the name
    is_first_line = True
    in_list = False
    
    # Parse the markdown content and add it to the Word document
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            continue
            
        # Detect section headings (bold text at start of line)
        if line.startswith('**') and line.endswith('**') and not line.startswith('* **'):
            # This is a section heading
            heading_text = line.strip('**')
            if is_first_line:
                # First line is the name - make it Heading 1
                doc.add_heading(heading_text, level=1)
                is_first_line = False
            else:
                # Other section headings - make them Heading 2
                doc.add_heading(heading_text, level=2)
        
        # Handle lists (bulleted and numbered)
        elif line.startswith(('*', '1.', '   *')):
            if line.startswith('*   '):
                text = process_markdown_formatting(line[4:].strip())
                doc.add_paragraph(text, style='List Bullet')
            elif line.startswith('* '):
                text = process_markdown_formatting(line[2:].strip())
                doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'\d+\. ', line):
                text = process_markdown_formatting(line[line.index('.') + 1:].strip())
                doc.add_paragraph(text, style='List Number')
            elif line.startswith('    *   ') or line.startswith('    * '):
                text = process_markdown_formatting(line.strip('    * ').strip())
                p = doc.add_paragraph(text, style='List Bullet 2')
            in_list = True

        # Handle code blocks
        elif line.startswith('```'):
            # Extract code language if specified
            language = line[3:].strip() or None
            
            # Read until the end of the code block
            code_block = []
            for code_line in lines[lines.index(line) + 1:]:
                if code_line.strip() == '```':
                    break
                code_block.append(code_line.rstrip())
            
            # Add the code block to the document
            paragraph = doc.add_paragraph('\n'.join(code_block), style='CodeBlock')
            
            # Skip the code block lines in the main loop
            lines = lines[:lines.index(line) + 1 + len(code_block) + 1] + lines[lines.index(line) + 1 + len(code_block) + 1:]
            
            continue
        
        # Handle images
        elif line.startswith('!['):
            # Extract alt text and image path
            alt_text = line[line.index('[') + 1:line.index(']')]
            image_path = line[line.index('(') + 1:line.index(')')]
            
            # Add the image to the document
            try:
                doc.add_picture(image_path, width=Inches(6.0))  # Adjust width as needed
            except Exception as e:
                print(f"Error adding image {image_path}: {e}")
                paragraph = doc.add_paragraph(f"Image: {alt_text} - {image_path} (Error)")
            
            continue
        
        # Handle tables
        elif '|' in line:
            # Detect the table header and separator
            if lines.index(line) + 1 < len(lines) and '---' in lines[lines.index(line) + 1]:
                # Extract the table headers
                header_line = line.strip()
                headers = [h.strip() for h in header_line.split('|') if h.strip()]
                
                # Extract the table data
                data_lines = []
                for table_line in lines[lines.index(line) + 2:]:
                    if '|' not in table_line:
                        break
                    data_lines.append(table_line.strip())
                
                # Create the table in the document
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Add the headers to the table
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                
                # Add the data to the table
                for data_line in data_lines:
                    data = [d.strip() for d in data_line.split('|') if d.strip()]
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(data):
                        if i < len(row_cells):
                            row_cells[i].text = cell_data
                
                # Skip the table lines in the main loop
                lines = lines[:lines.index(line) + 1 + len(data_lines) + 1] + lines[lines.index(line) + 1 + len(data_lines) + 1:]
                
                continue
        
        # Handle markdown links
        elif '[' in line and '](' in line and ')' in line:
            # Extract all links in the line
            current_pos = 0
            paragraph = doc.add_paragraph()
            
            while current_pos < len(line):
                link_start = line.find('[', current_pos)
                
                # If no more links, add the rest of the text
                if link_start == -1:
                    if current_pos < len(line):
                        remaining_text = line[current_pos:]
                        # Process any markdown formatting in the remaining text
                        add_formatted_text(paragraph, remaining_text)
                    break
                    
                # Add text before the link
                if link_start > current_pos:
                    before_link_text = line[current_pos:link_start]
                    # Process any markdown formatting in text before the link
                    add_formatted_text(paragraph, before_link_text)
                
                # Extract link text and URL
                text_start = link_start + 1
                text_end = line.find(']', text_start)
                url_start = text_end + 2  # Skip ']('
                url_end = line.find(')', url_start)
                
                if text_end == -1 or url_end == -1:
                    # Malformed link, just add as text
                    remaining_text = line[current_pos:]
                    # Process any markdown formatting in the remaining text
                    add_formatted_text(paragraph, remaining_text)
                    break
                
                link_text = line[text_start:text_end]
                url = line[url_start:url_end]
                
                # Add link as blue, underlined text
                run = paragraph.add_run(link_text)
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
                run.underline = True
                
                # Move current position past this link
                current_pos = url_end + 1
            
            in_list = False
        else:
            # Regular paragraph with possible markdown formatting
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, line)
            in_list = False

    # Save the Word document
    doc.save(output_file_path)
    print(f"Document successfully created: {output_file_path}")

def process_markdown_formatting(text):
    """
    Process markdown formatting in text before creating a paragraph.
    Use this for simple formatting in texts that will be directly added as paragraphs.
    """
    # This function is used when we can't add runs to an existing paragraph (like list items)
    # For complex formatting, use add_formatted_text instead
    
    # Remove markdown bold formatting as we can't apply it in this simple mode
    # We'll rely on the actual Word styling for lists
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    return text

def add_formatted_text(paragraph, text):
    """
    Add text to a paragraph with markdown formatting applied.
    """
    # Process bold text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]  # Remove the ** markers
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Regular text
            if part:
                paragraph.add_run(part)

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python script.py <input_md_file> <output_docx_file>")
    else:
        input_md_file = sys.argv[1]
        output_docx_file = sys.argv[2]
        generate_word_from_md(input_md_file, output_docx_file)